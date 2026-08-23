#!/usr/bin/env python3
"""Build deterministic retail price-accuracy evidence and mutation packs.

The source scenarios are synthetic.  They share an exact document schema so a
benchmark can change facts and jurisdiction without changing the file-reading
problem.  Legal rows are a triage register, not legal advice: every one is
flagged for current attorney validation, and only specifically identified rows
claim primary-source triage.
"""
from __future__ import annotations

if not __debug__:
    raise RuntimeError("validation requires assertions; Python optimization is unsupported")

import argparse
from datetime import datetime
import hashlib
import json
from pathlib import Path
import re
import sys
import tempfile
from typing import Any
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = ROOT / "research" / "retail-price-accuracy"
CONTRACT = ROOT / "mcp" / "v4" / "contracts" / "retail-compliance.json"
# Committed research input (not a build output): AI-assisted primary-source
# triage of the 45 research-queue jurisdictions.  Its rows never overwrite the
# frozen rc_jurisdiction_rules seed table; they feed only the separate
# register-v2 evidence pack and remain pending attorney validation.
RESEARCH_V2 = OUTPUT_ROOT / "jurisdiction-research-v2.json"
BUILD_DATE = "2026-08-22"
FIXED_DATETIME = datetime(2026, 8, 22, 12, 0, 0)
FIXED_ZIP_TIME = (2026, 8, 22, 12, 0, 0)
NIST_URL = "https://www.nist.gov/pml/owm/nist-handbook-130-current-edition"

BLUE = "17365D"
PALE_BLUE = "D9EAF7"
GOLD = "D7A62A"
PALE_GOLD = "FFF3CD"
RED = "A61B1B"
PALE_RED = "FCE8E6"
GREEN = "1F6B45"
PALE_GREEN = "E2F0D9"
GRAY = "E7E6E6"


JURISDICTIONS = [
    ("AL", "Alabama", "https://legislature.state.al.us/aliswww/ISD/CodeOfAlabama/1975/coatoc.htm"),
    ("AK", "Alaska", "https://www.akleg.gov/basis/statutes.asp"),
    ("AZ", "Arizona", "https://www.azleg.gov/arstitle/"),
    ("AR", "Arkansas", "https://www.arkleg.state.ar.us/ArkansasLaw"),
    ("CA", "California", "https://leginfo.legislature.ca.gov/faces/codes.xhtml"),
    ("CO", "Colorado", "https://leg.colorado.gov/agencies/office-legislative-legal-services/colorado-revised-statutes"),
    ("CT", "Connecticut", "https://www.cga.ct.gov/current/pub/titles.htm"),
    ("DE", "Delaware", "https://delcode.delaware.gov/"),
    ("FL", "Florida", "https://www.leg.state.fl.us/Statutes/"),
    ("GA", "Georgia", "https://www.legis.ga.gov/laws/georgia-code"),
    ("HI", "Hawaii", "https://www.capitol.hawaii.gov/hrscurrent/"),
    ("ID", "Idaho", "https://legislature.idaho.gov/statutesrules/idstat/"),
    ("IL", "Illinois", "https://www.ilga.gov/legislation/ilcs/ilcs.asp"),
    ("IN", "Indiana", "https://iga.in.gov/laws/2025/ic/titles/1"),
    ("IA", "Iowa", "https://www.legis.iowa.gov/law/iowaCode"),
    ("KS", "Kansas", "https://www.ksrevisor.gov/statutes/ksa_chapters.html"),
    ("KY", "Kentucky", "https://apps.legislature.ky.gov/law/statutes/"),
    ("LA", "Louisiana", "https://legis.la.gov/Legis/Laws_Toc.aspx"),
    ("ME", "Maine", "https://legislature.maine.gov/statutes/"),
    ("MD", "Maryland", "https://mgaleg.maryland.gov/mgawebsite/Laws/Statutes"),
    ("MA", "Massachusetts", "https://malegislature.gov/Laws/GeneralLaws"),
    ("MI", "Michigan", "https://www.legislature.mi.gov/Laws/MCL"),
    ("MN", "Minnesota", "https://www.revisor.mn.gov/statutes/"),
    ("MS", "Mississippi", "https://www.legislature.ms.gov/laws/mississippi-code/"),
    ("MO", "Missouri", "https://revisor.mo.gov/main/Home.aspx"),
    ("MT", "Montana", "https://leg.mt.gov/bills/mca/index.html"),
    ("NE", "Nebraska", "https://nebraskalegislature.gov/laws/browse-statutes.php"),
    ("NV", "Nevada", "https://www.leg.state.nv.us/NRS/"),
    ("NH", "New Hampshire", "https://gc.nh.gov/rsa/html/nhtoc.htm"),
    ("NJ", "New Jersey", "https://www.njleg.state.nj.us/new-jersey-statutes"),
    ("NM", "New Mexico", "https://nmonesource.com/nmos/en/nav.do"),
    ("NY", "New York", "https://www.nysenate.gov/legislation/laws/CONSOLIDATED"),
    ("NC", "North Carolina", "https://www.ncleg.gov/Laws/GeneralStatutes"),
    ("ND", "North Dakota", "https://ndlegis.gov/general-information/north-dakota-century-code"),
    ("OH", "Ohio", "https://codes.ohio.gov/ohio-revised-code"),
    ("OK", "Oklahoma", "https://oksenate.gov/publications/oklahoma-statutes"),
    ("OR", "Oregon", "https://www.oregonlegislature.gov/bills_laws/Pages/ORS.aspx"),
    ("PA", "Pennsylvania", "https://www.legis.state.pa.us/WU01/LI/LI/CT/HTM/00/00.HTM"),
    ("RI", "Rhode Island", "https://webserver.rilegislature.gov/Statutes/"),
    ("SC", "South Carolina", "https://www.scstatehouse.gov/code/statmast.php"),
    ("SD", "South Dakota", "https://sdlegislature.gov/Statutes"),
    ("TN", "Tennessee", "https://www.capitol.tn.gov/Archives/Joint/publications/Title%20and%20Chapter%20list.htm"),
    ("TX", "Texas", "https://statutes.capitol.texas.gov/"),
    ("UT", "Utah", "https://le.utah.gov/xcode/code.html"),
    ("VT", "Vermont", "https://legislature.vermont.gov/statutes/"),
    ("VA", "Virginia", "https://law.lis.virginia.gov/vacode"),
    ("WA", "Washington", "https://app.leg.wa.gov/rcw/"),
    ("WV", "West Virginia", "https://code.wvlegislature.gov/"),
    ("WI", "Wisconsin", "https://docs.legis.wisconsin.gov/statutes/statutes"),
    ("WY", "Wyoming", "https://wyoleg.gov/statutes/compress/title40.pdf"),
    ("DC", "District of Columbia", "https://code.dccouncil.gov/us/dc/council/code"),
]


PRIMARY_RULES: dict[str, dict[str, Any]] = {
    "CA": {
        "rule_tier": "enhanced_lowest_price",
        "price_standard": "Charge no more than the lowest advertised, posted, marked, displayed, or quoted price, subject to statutory exceptions.",
        "consumer_remedy": "Prompt difference refund; preserve statutory rights; escalate systemic or weight/measure issues.",
        "authority": "Cal. Bus. & Prof. Code §§ 12024.2, 13350",
        "source_url": "https://leginfo.legislature.ca.gov/faces/codes_displaySection.xhtml?lawCode=BPC&sectionNum=12024.2",
        "notice_window_days": None,
        "payment_window_days": None,
    },
    "MI": {
        "rule_tier": "enhanced_scanner_bonus",
        "price_standard": "Clearly display price and do not charge more through an automatic checkout system.",
        "consumer_remedy": "After a completed transaction and qualifying notice, refund the difference and apply validated statutory bonus mechanics.",
        "authority": "Shopping Reform and Modernization Act, 2011 PA 15",
        "source_url": "https://www.legislature.mi.gov/documents/mcl/pdf/mcl-Act-15-of-2011.pdf",
        "notice_window_days": 30,
        "payment_window_days": 2,
    },
    "DC": {
        "rule_tier": "udap_and_unit_price",
        "price_standard": "Avoid material price misrepresentation and display unit pricing clearly and non-deceptively where applicable.",
        "consumer_remedy": "Refund and legal escalation under applicable consumer-protection procedures; no fixed bonus assumed.",
        "authority": "D.C. Code §§ 28-3904, 28-5207",
        "source_url": "https://code.dccouncil.gov/us/dc/council/code/sections/28-3904",
        "notice_window_days": None,
        "payment_window_days": None,
    },
    "NY": {
        "rule_tier": "enhanced_item_pricing",
        "price_standard": "Use the lowest advertised, written, posted, or marked price where Agriculture and Markets Law § 197-b applies.",
        "consumer_remedy": "Post and follow the applicable refund policy; apply local inspection and penalty rules after validation.",
        "authority": "N.Y. Agric. & Mkts. Law § 197-b",
        "source_url": "https://www.nysenate.gov/legislation/laws/AGM/197-B",
        "notice_window_days": None,
        "payment_window_days": None,
    },
    "MA": {
        "rule_tier": "enhanced_price_disclosure",
        "price_standard": "Apply the correct price and receipt/disclosure requirements under 202 CMR 7.00 where applicable.",
        "consumer_remedy": "Apply validated price-accuracy remedy and disclosure rules without limiting statutory rights.",
        "authority": "202 CMR 7.00",
        "source_url": "https://www.mass.gov/regulations/202-CMR-700-price-disclosure",
        "notice_window_days": None,
        "payment_window_days": None,
    },
    "CT": {
        "rule_tier": "enhanced_scanner_accuracy",
        "price_standard": "Use the posted price and meet applicable scanner-accuracy and consumer-remedy rules.",
        "consumer_remedy": "Apply the point-of-sale overcharge remedy only after confirming current statutory conditions and exclusions.",
        "authority": "Conn. Gen. Stat. ch. 417",
        "source_url": "https://www.cga.ct.gov/current/pub/chap_417.htm",
        "notice_window_days": None,
        "payment_window_days": None,
    },
}


SCENARIOS = [
    {
        "slug": "scenarios/ca-price-weight-duplicate-scan-v1",
        "scenario_id": "CA-PRICE-001",
        "jurisdiction_code": "CA",
        "jurisdiction_name": "California",
        "store_id": "CA-1842",
        "incident_number": "RG-2026-0042",
        "transaction_id": "TX-CA-000184",
        "lane": "SCO-07",
        "customer": "Customer CA-184",
        "occurred_at": "2026-08-14T16:42:00Z",
        "root_cause": "A retry event was treated as a new scan while the catalog cache retained an expired shelf price.",
        "displayed": 4.99,
        "charged": 5.39,
        "duplicate_price": 5.49,
        "weight_label": 3.00,
        "weight_charged": 3.18,
        "tax": 3.83,
    },
    {
        "slug": "mutations/mi-price-weight-duplicate-scan-v1",
        "scenario_id": "MI-PRICE-001",
        "jurisdiction_code": "MI",
        "jurisdiction_name": "Michigan",
        "store_id": "MI-0907",
        "incident_number": "RG-2026-0051",
        "transaction_id": "TX-MI-000291",
        "lane": "SCO-12",
        "customer": "Customer MI-291",
        "occurred_at": "2026-08-16T18:07:00Z",
        "root_cause": "A barcode debounce rule was disabled after a terminal update and the promotional price feed missed one SKU.",
        "displayed": 3.79,
        "charged": 4.29,
        "duplicate_price": 6.25,
        "weight_label": 2.00,
        "weight_charged": 2.16,
        "tax": 2.94,
    },
    {
        "slug": "mutations/dc-price-weight-duplicate-scan-v1",
        "scenario_id": "DC-PRICE-001",
        "jurisdiction_code": "DC",
        "jurisdiction_name": "District of Columbia",
        "store_id": "DC-0311",
        "incident_number": "RG-2026-0064",
        "transaction_id": "TX-DC-000338",
        "lane": "SCO-04",
        "customer": "Customer DC-338",
        "occurred_at": "2026-08-18T13:23:00Z",
        "root_cause": "A unit-price label mapped to a stale item record and a rapid rescan was not challenged.",
        "displayed": 7.49,
        "charged": 7.99,
        "duplicate_price": 4.75,
        "weight_label": 1.50,
        "weight_charged": 1.64,
        "tax": 4.12,
    },
]




def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_ooxml(path: Path) -> None:
    """Remove wall-clock ZIP metadata while preserving OOXML member bytes."""
    temporary = path.with_suffix(path.suffix + ".normalized")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as destination:
        for original in source.infolist():
            member = zipfile.ZipInfo(original.filename, FIXED_ZIP_TIME)
            member.compress_type = original.compress_type
            member.comment = original.comment
            member.create_system = original.create_system
            member.external_attr = original.external_attr
            member.internal_attr = original.internal_attr
            member.flag_bits = original.flag_bits
            data = source.read(original.filename)
            if original.filename == "docProps/core.xml":
                for tag in (b"created", b"modified"):
                    pattern = rb"(<dcterms:" + tag + rb"[^>]*>)[^<]*(</dcterms:" + tag + rb">)"
                    data = re.sub(pattern, rb"\g<1>2026-08-22T12:00:00Z\g<2>", data)
            destination.writestr(member, data)
    temporary.replace(path)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell_text(cell, text: Any, *, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_doc(document: Document, title: str, scenario: dict[str, Any]) -> None:
    document.core_properties.created = FIXED_DATETIME
    document.core_properties.modified = FIXED_DATETIME
    document.core_properties.revision = 1
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in (("Title", 22, BLUE), ("Heading 1", 15, BLUE), ("Heading 2", 11, GREEN)):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.text = f"RETAILGUARD / SYNTHETIC EVIDENCE / {scenario['scenario_id']}"
    header.style = document.styles["Caption"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Simulation only • synthetic parties and transactions • attorney validation required")
    document.add_heading(title, 0)
    subtitle = document.add_paragraph()
    subtitle.add_run(f"{scenario['jurisdiction_name']} • {scenario['incident_number']} • Built {BUILD_DATE}").bold = True
    disclaimer = document.add_paragraph("SIMULATION ONLY. This record is synthetic and is not legal advice. Legal conclusions and remedies require current, jurisdiction-specific attorney review.")
    disclaimer.style = document.styles["Quote"]


def add_kv_table(document: Document, values: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in values:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=BLUE)
        set_cell_shading(cells[0], PALE_BLUE)
        set_cell_text(cells[1], value)


def build_incident_doc(path: Path, scenario: dict[str, Any]) -> None:
    document = Document()
    configure_doc(document, "Price Accuracy Incident Report", scenario)
    add_kv_table(document, [
        ("Incident", scenario["incident_number"]),
        ("Store / lane", f"{scenario['store_id']} / {scenario['lane']}"),
        ("Transaction", scenario["transaction_id"]),
        ("Timestamp", scenario["occurred_at"]),
        ("Channel", "Self-checkout"),
        ("Status", "Open — legal and control review"),
    ])
    document.add_heading("Executive finding", level=1)
    document.add_paragraph(
        f"The evidence indicates three distinct price-integrity failures: a probable duplicate scan of "
        f"${scenario['duplicate_price']:.2f}, a shelf-to-register mismatch of "
        f"${scenario['charged'] - scenario['displayed']:.2f}, and a packaged-weight variance from "
        f"{scenario['weight_label']:.2f} lb labeled to {scenario['weight_charged']:.2f} lb charged. "
        "The duplicate charge must not be described as a confirmed system defect until terminal logs and video are preserved and reviewed."
    )
    document.add_heading("Preservation and investigation", level=1)
    for text in [
        "Preserve the receipt, terminal event stream, catalog and promotional-price snapshots, scale calibration logs, shelf-label photographs, refund contacts, relevant video, change tickets, and audit results.",
        "Separate duplicate-scan, advertised-price, and net-weight theories; calculate transaction-level amounts without extrapolating a class until the population query is validated.",
        "Apply a litigation hold to the incident window and document custodians; record chain of custody and hashes for exported evidence.",
        "Escalate any systemic pattern, regulator contact, threatened aggregate claim, or statutory-bonus question to qualified counsel.",
    ]:
        document.add_paragraph(text, style="List Bullet")
    document.add_heading("Root-cause hypothesis", level=1)
    document.add_paragraph(scenario["root_cause"])
    document.add_heading("Immediate customer remediation", level=1)
    document.add_paragraph(
        "Promptly review the customer's evidence, refund every verified overcharge through the original or otherwise lawful payment method, separately apply any validated jurisdiction-specific remedy, and issue a notice that does not waive or limit statutory rights. Do not require a release for an ordinary price correction."
    )
    document.add_heading("Legal triage", level=1)
    rule = PRIMARY_RULES[scenario["jurisdiction_code"]]
    add_kv_table(document, [
        ("Research tier", rule["rule_tier"]),
        ("Primary authority", rule["authority"]),
        ("Triage standard", rule["price_standard"]),
        ("Validation flag", "Attorney validation required before implementation"),
    ])
    document.add_heading("Non-conclusion", level=1)
    document.add_paragraph(
        "Receipt language and contract terms cannot ensure that a retailer will not be sued. The control objective is accurate charging, rapid detection, complete remediation, preserved evidence, and non-misleading communication."
    )
    document.save(path)
    normalize_ooxml(path)


def build_policy_doc(path: Path, scenario: dict[str, Any]) -> None:
    document = Document()
    configure_doc(document, "Customer Price Accuracy & Self-Checkout Policy", scenario)
    sections = [
        ("1. Purpose and scope", "This policy governs displayed prices, item scans, weights, receipts, refunds, audit evidence, and customer communications for staffed and self-checkout sales."),
        ("2. Lowest supported price control", "At checkout, the retailer will apply the lowest price supported by a current advertisement, shelf label, item marking, display, or other controlling source when required by applicable law. Disputed exceptions are escalated; they are not resolved by hiding or disclaiming the displayed price."),
        ("3. Duplicate-scan control", "A probable rapid duplicate scan must trigger an on-screen confirmation or associate review. Quantity changes remain available, but the event and shopper choice must be logged."),
        ("4. Weight and measure control", "Packaged and random-weight goods must be sold using validated scales, current tare, and accurate labels. A variance alert pauses the sale and triggers scale, label, and catalog review."),
        ("5. Receipt and checkout notice", "Receipts must identify each item, quantity or weight, unit or extended price, discounts, tax, total, store, lane, timestamp, and a reachable price-accuracy channel. The notice must invite review without shifting the retailer's compliance duty to the customer."),
        ("6. Refunds and additional remedies", "Refund every verified overcharge promptly. Apply jurisdiction-specific bonuses, free-item remedies, notices, or deadlines only after current legal validation. Never use a refund receipt, clickwrap term, or customer-service script to waive nonwaivable rights."),
        ("7. Evidence and escalation", "Preserve event logs, receipts, catalog snapshots, shelf evidence, weights-and-measures records, refunds, complaints, control versions, and audit results. Escalate systemic patterns, regulator inquiries, and aggregate-claim threats to counsel."),
        ("8. Testing and governance", "Run risk-based store samples informed by NIST Handbook 130 EPPV while applying controlling state and local law. Record the sample population, exceptions, correction, owner, and retest."),
    ]
    for heading, body in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(body)
    document.add_heading("Approved customer-facing wording", level=1)
    wording = document.add_table(rows=1, cols=2)
    wording.style = "Table Grid"
    wording.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_text(wording.rows[0].cells[0], "Touchpoint", bold=True, color="FFFFFF")
    set_cell_text(wording.rows[0].cells[1], "Text", bold=True, color="FFFFFF")
    for cell in wording.rows[0].cells:
        set_cell_shading(cell, BLUE)
    rows = [
        ("Checkout", "Please review item, quantity, weight, price, discounts, and total before payment. Ask an associate about any difference."),
        ("Receipt", "If a price or scan appears incorrect, keep your receipt and contact an associate or the price-accuracy desk for prompt review and any amount due."),
        ("Savings clause", "This notice does not limit any right or remedy provided by applicable law."),
    ]
    for label, body in rows:
        cells = wording.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], body)
    document.add_heading("Jurisdiction deployment gate", level=1)
    document.add_paragraph(
        "No jurisdiction-specific variant may move beyond legal review unless the source, effective date, scope, exclusions, remedy, notice period, payment period, receipt requirement, private right, preemption/local overlay, and operational owner have been validated by qualified counsel."
    )
    document.save(path)
    normalize_ooxml(path)


def style_sheet(sheet, *, freeze: str = "A2", filter_row: int = 1) -> None:
    sheet.freeze_panes = freeze
    sheet.auto_filter.ref = f"A{filter_row}:{get_column_letter(sheet.max_column)}{sheet.max_row}"
    for cell in sheet[filter_row]:
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    sheet.row_dimensions[filter_row].height = 34
    for column in range(1, sheet.max_column + 1):
        values = [str(sheet.cell(row=row, column=column).value or "") for row in range(1, min(sheet.max_row, 80) + 1)]
        width = min(48, max(11, max(map(len, values), default=10) + 2))
        sheet.column_dimensions[get_column_letter(column)].width = width
    for row in sheet.iter_rows(min_row=filter_row + 1):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def build_event_log(path: Path, scenario: dict[str, Any]) -> None:
    workbook = Workbook()
    workbook.properties.created = FIXED_DATETIME
    workbook.properties.modified = FIXED_DATETIME
    events = workbook.active
    events.title = "Events"
    events.append(["event_id", "timestamp_utc", "transaction_id", "lane", "event_type", "sku", "value", "operator", "evidence_status"])
    events_data = [
        (1, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "basket_open", "", "", "shopper", "preserved"),
        (2, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "scan", "04120-88214", scenario["duplicate_price"], "shopper", "preserved"),
        (3, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "scan_retry", "04120-88214", scenario["duplicate_price"], "shopper", "preserved"),
        (4, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "catalog_price", "00854-10219", scenario["charged"], "system", "preserved"),
        (5, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "shelf_snapshot", "00854-10219", scenario["displayed"], "associate", "preserved"),
        (6, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "scale_weight", "02941-55003", scenario["weight_charged"], "system", "preserved"),
        (7, scenario["occurred_at"], scenario["transaction_id"], scenario["lane"], "tender_complete", "", "", "shopper", "preserved"),
    ]
    for row in events_data:
        events.append(row)
    style_sheet(events)

    lines = workbook.create_sheet("Line Items")
    lines.append(["line_id", "transaction_id", "sku", "description", "quantity", "displayed_unit_price", "charged_unit_price", "label_weight_lb", "charged_weight_lb", "extended_price", "computed_overcharge", "anomaly_type"])
    line_data = [
        (1, scenario["transaction_id"], "04120-88214", "Organic oranges bag", 2, scenario["duplicate_price"], scenario["duplicate_price"], scenario["weight_label"], scenario["weight_label"], scenario["duplicate_price"] * 2, None, "duplicate_scan"),
        (2, scenario["transaction_id"], "00854-10219", "Sparkling water", 1, scenario["displayed"], scenario["charged"], None, None, scenario["charged"], None, "shelf_register_mismatch"),
        (3, scenario["transaction_id"], "02941-55003", "Random-weight produce", 1, 2.99, 2.99, scenario["weight_label"], scenario["weight_charged"], round(2.99 * scenario["weight_charged"], 2), None, "weight_variance"),
    ]
    for row_index, row in enumerate(line_data, 2):
        lines.append(row[:-2] + (f"=ROUND(MAX(0,(G{row_index}-F{row_index})*E{row_index})+IF(L{row_index}=\"duplicate_scan\",F{row_index},0)+IF(L{row_index}=\"weight_variance\",MAX(0,(I{row_index}-H{row_index})*G{row_index}),0),2)", row[-1]))
    style_sheet(lines)
    lines.conditional_formatting.add(f"K2:K{lines.max_row}", CellIsRule(operator="greaterThan", formula=["0"], fill=PatternFill("solid", fgColor=PALE_RED)))
    for column in (6, 7, 10, 11):
        for cell in list(lines.columns)[column - 1][1:]:
            cell.number_format = '$0.00'

    controls = workbook.create_sheet("Control Tests")
    controls.append(["control_code", "test", "expected", "observed", "result", "owner", "required_action"])
    for row in [
        ("SCO-DUP-01", "Rapid same-SKU scan challenge", "confirmation or associate review", "no challenge", "FAIL", "POS Engineering", "enable and regression-test debounce control"),
        ("PRICE-SYNC-02", "Shelf/register price comparison", "honor lower supported price", "stale catalog value charged", "FAIL", "Pricing Operations", "quarantine SKU and reconcile feed"),
        ("WEIGHT-03", "Label/charged weight comparison", "within validated tolerance", "charged weight exceeds label", "FAIL", "Quality", "hold sale and verify scale, tare, and label"),
    ]:
        controls.append(row)
    style_sheet(controls)
    for cell in controls["E"][1:]:
        cell.fill = PatternFill("solid", fgColor=PALE_RED if cell.value == "FAIL" else PALE_GREEN)
        cell.font = Font(color=RED if cell.value == "FAIL" else GREEN, bold=True)

    dictionary = workbook.create_sheet("Data Dictionary")
    dictionary.append(["field", "meaning"])
    for row in [
        ("computed_overcharge", "Synthetic issue-spotting calculation; validate refund amount and remedy under applicable law."),
        ("evidence_status", "Whether an immutable export and chain-of-custody record exist."),
        ("anomaly_type", "duplicate_scan, shelf_register_mismatch, or weight_variance."),
        ("SIMULATION", "All parties, stores, transactions, and evidence in this workbook are synthetic."),
    ]:
        dictionary.append(row)
    style_sheet(dictionary)
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.save(path)
    normalize_ooxml(path)


def jurisdiction_rows() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for index, (code, name, portal) in enumerate(JURISDICTIONS, 1):
        primary = PRIMARY_RULES.get(code)
        if primary:
            row = {
                "id": index,
                "jurisdiction_code": code,
                "jurisdiction_name": name,
                "rule_tier": primary["rule_tier"],
                "research_status": "primary_source_triaged",
                "price_standard": primary["price_standard"],
                "consumer_remedy": primary["consumer_remedy"],
                "notice_window_days": primary["notice_window_days"],
                "payment_window_days": primary["payment_window_days"],
                "authority": primary["authority"],
                "source_url": primary["source_url"],
                "last_verified": BUILD_DATE,
                "verification_note": "Official text triaged for benchmark design; confirm current text, scope, local overlays, remedies, and effective date.",
                "attorney_validation_required": 1,
            }
        else:
            row = {
                "id": index,
                "jurisdiction_code": code,
                "jurisdiction_name": name,
                "rule_tier": "baseline_research_pending",
                "research_status": "official_portal_identified_not_substantively_validated",
                "price_standard": "Operational floor: prevent duplicate scans, reconcile displayed and charged prices, use accurate weights, and preserve evidence; controlling law not yet coded.",
                "consumer_remedy": "Promptly refund verified overcharge and escalate jurisdiction-specific remedy analysis; do not assume a bonus, free-item remedy, waiver, or notice deadline.",
                "notice_window_days": None,
                "payment_window_days": None,
                "authority": "Research queue: state consumer-protection, pricing, and weights-and-measures law; NIST HB 130 is an inspection baseline, not the law.",
                "source_url": portal,
                "last_verified": BUILD_DATE,
                "verification_note": "Official code portal identified only. Substantive rule, remedies, local overlays, and current effective text remain attorney work.",
                "attorney_validation_required": 1,
            }
        rows.append(row)
    assert len(rows) == 51 and len({row["jurisdiction_code"] for row in rows}) == 51
    return rows


def jurisdiction_research_v2_rows() -> list[dict[str, Any]]:
    """Load the committed authority map without promoting it into legal rules."""
    payload = json.loads(RESEARCH_V2.read_text("utf-8"))
    defaults = payload.get("defaults") or {}
    source_rows = payload.get("jurisdictions") or []
    rows = [{**defaults, **row} for row in source_rows]
    expected_codes = {code for code, _name, _portal in JURISDICTIONS}
    actual_codes = {row.get("code") for row in rows}
    if payload.get("schema_version") != 2 or len(rows) != 51 or actual_codes != expected_codes:
        raise RuntimeError("jurisdiction research v2 must map exactly 50 states plus District of Columbia")
    required = {
        "code", "name", "citation", "authority_url", "source_kind", "authority_focus",
        "mapping_status", "operational_baseline", "substantive_legal_opinion",
        "private_remedy_encoded", "current_text_and_local_overlays_validated",
        "attorney_validation_required",
    }
    for row in rows:
        missing = sorted(required - row.keys())
        if missing:
            raise RuntimeError(f"jurisdiction research v2 {row.get('code')} missing {missing}")
        if not str(row["authority_url"]).startswith("https://"):
            raise RuntimeError(f"jurisdiction research v2 {row['code']} must use an HTTPS authority URL")
        if row["substantive_legal_opinion"] or row["private_remedy_encoded"]:
            raise RuntimeError(f"jurisdiction research v2 {row['code']} improperly encodes legal conclusions")
        if row["current_text_and_local_overlays_validated"] or not row["attorney_validation_required"]:
            raise RuntimeError(f"jurisdiction research v2 {row['code']} bypasses the attorney gate")
    order = {code: index for index, (code, _name, _portal) in enumerate(JURISDICTIONS)}
    return sorted(rows, key=lambda row: order[row["code"]])


def build_jurisdiction_register(path: Path, scenario: dict[str, Any]) -> None:
    workbook = Workbook()
    workbook.properties.created = FIXED_DATETIME
    workbook.properties.modified = FIXED_DATETIME
    sheet = workbook.active
    sheet.title = "Jurisdiction Register"
    headers = [
        "jurisdiction_code", "jurisdiction_name", "rule_tier", "research_status",
        "price_standard", "consumer_remedy", "notice_window_days", "payment_window_days",
        "authority", "source_url", "last_verified", "verification_note",
        "attorney_validation_required", "scenario_relevance",
    ]
    sheet.append(headers)
    for row in jurisdiction_rows():
        sheet.append([row.get(header) for header in headers[:-1]] + [
            "primary scenario" if row["jurisdiction_code"] == scenario["jurisdiction_code"] else "comparison queue"
        ])
    style_sheet(sheet)
    sheet.auto_filter.ref = f"A1:N{sheet.max_row}"
    for row in range(2, sheet.max_row + 1):
        status = sheet.cell(row, 4).value
        fill = PALE_GREEN if status == "primary_source_triaged" else PALE_GOLD
        sheet.cell(row, 4).fill = PatternFill("solid", fgColor=fill)
    sources = workbook.create_sheet("Primary Sources")
    sources.append(["source_id", "type", "authority_or_matter", "url", "use", "limitation"])
    source_rows = [
        ("FED-NIST-HB130", "federal measurement handbook", "NIST Handbook 130 / EPPV", NIST_URL, "inspection-method baseline", "not itself state law"),
        ("WMT-RECTOR", "official federal opinion", "Rector v. Walmart, D.D.C. 1:24-cv-00658-RC", "https://www.govinfo.gov/content/pkg/USCOURTS-dcd-1_24-cv-00658/pdf/USCOURTS-dcd-1_24-cv-00658-1.pdf", "shelf/register allegation and first-filed procedure", "not a merits ruling or settlement"),
        ("WMT-KAHN", "official federal opinion", "Kahn v. Walmart, 107 F.4th 585 (7th Cir. 2024)", "https://media.ca7.uscourts.gov/cgi-bin/OpinionsWeb/processWebInputExternal.pl?Path=Y2024%2FD07-03%2FC%3A23-1751%3AJ%3AHamilton%3Aaut%3AT%3AfnOp%3AN%3A3231141%3AS%3A0&Submit=Display", "scanner-price pleading and consumer-protection analysis", "case-specific procedural opinion"),
        ("WMT-WEIGHTED", "official court settlement record", "Kukorinis v. Walmart, M.D. Fla. 8:22-cv-02402", "https://ecf.flmd.uscourts.gov/cgi-bin/show_public_doc?2022-02402-132-8-cv=", "weighted-goods settlement", "Florida federal case; Walmart denied wrongdoing"),
        ("CA-AG-2012", "official enforcement release", "California checkout overcharge judgment enforcement", "https://oag.ca.gov/node/30977", "controls and customer remediation", "enforcement summary, not statutory text"),
        ("CA-DA-2025", "official enforcement release", "California price/weight settlement", "https://da.santaclaracounty.gov/walmart-overcharged-customers-will-pay-56-million-settle-consumer-protection-lawsuit", "price and weight compliance", "enforcement summary; allegations resolved by settlement"),
    ]
    for row in source_rows:
        sources.append(row)
    style_sheet(sources)
    method = workbook.create_sheet("Methodology")
    method.append(["field", "method"])
    for row in [
        ("Scope", "50 states plus District of Columbia; 51 rows total."),
        ("Validated subset", "CA, MI, DC, NY, MA, CT received benchmark-level primary-source triage."),
        ("Unvalidated subset", "Official code portal identified; substantive rule intentionally not asserted."),
        ("Operational baseline", "NIST HB 130 EPPV informs inspection design but is not substituted for governing law."),
        ("Deployment gate", "Current qualified-counsel review of state and local law, applicability, remedies, notices, effective dates, and preemption."),
        ("No guarantee", "Wording and controls reduce risk; they cannot ensure that litigation will not occur."),
    ]:
        method.append(row)
    style_sheet(method)
    workbook.save(path)
    normalize_ooxml(path)


def build_jurisdiction_register_v2(path: Path, scenario: dict[str, Any]) -> None:
    """Build a 51-jurisdiction issue-spotting register with an explicit legal gate."""
    payload = json.loads(RESEARCH_V2.read_text("utf-8"))
    rows = jurisdiction_research_v2_rows()
    workbook = Workbook()
    workbook.properties.created = FIXED_DATETIME
    workbook.properties.modified = FIXED_DATETIME
    sheet = workbook.active
    sheet.title = "Authority Map"
    headers = [
        "jurisdiction_code", "jurisdiction_name", "citation", "authority_focus",
        "authority_url", "source_kind", "mapping_status", "operational_baseline",
        "substantive_legal_opinion", "private_remedy_encoded",
        "current_text_and_local_overlays_validated", "attorney_validation_required",
        "scenario_relevance",
    ]
    sheet.append(headers)
    for row in rows:
        values = {
            "jurisdiction_code": row["code"],
            "jurisdiction_name": row["name"],
            **row,
            "scenario_relevance": (
                "primary scenario authority map"
                if row["code"] == scenario["jurisdiction_code"]
                else "national comparison authority map"
            ),
        }
        sheet.append([values[header] for header in headers])
    style_sheet(sheet)
    sheet.auto_filter.ref = f"A1:M{sheet.max_row}"
    for row_number in range(2, sheet.max_row + 1):
        sheet.cell(row_number, 7).fill = PatternFill("solid", fgColor=PALE_GOLD)
        sheet.cell(row_number, 12).fill = PatternFill("solid", fgColor=PALE_RED)

    controls = workbook.create_sheet("Control Baseline")
    controls.append(["control_id", "control", "verification_evidence", "legal_boundary"])
    control_rows = [
        ("POS-01", "Debounce or challenge rapid duplicate scans.", "immutable scan-event sequence and override log", "technical control; no legal conclusion"),
        ("POS-02", "Reconcile shelf, promotion, catalog, and checkout prices before sale.", "versioned price feed and mismatch alert", "apply controlling lowest-price rule only after counsel validation"),
        ("POS-03", "Use calibrated devices and reconcile label and charged weight.", "device certificate, raw weight event, label record", "local weights-and-measures rules may add duties"),
        ("POS-04", "Issue an intelligible itemized receipt and preserve transaction evidence.", "receipt, transaction export, content hashes", "state receipt requirements and retention periods require validation"),
        ("POS-05", "Promptly correct a verified overcharge without requiring a release.", "refund ledger and customer notice", "do not infer a bonus, free-item remedy, notice period, waiver, or release"),
        ("POS-06", "Escalate patterns, repeat failures, and jurisdiction-specific remedy questions.", "incident, audit, and counsel-approval records", "qualified counsel approves deployment"),
    ]
    for row in control_rows:
        controls.append(row)
    style_sheet(controls)

    methodology = workbook.create_sheet("Methodology")
    methodology.append(["field", "value"])
    method_rows = [
        ("as_of", payload["as_of"]),
        ("scope", payload["scope"]),
        ("national_index", payload["research_method"]["national_index"]),
        ("source_preference", " > ".join(payload["research_method"]["source_preference"])),
        ("mapped_rows", len(rows)),
        ("substantive_legal_opinions", 0),
        ("private_remedies_encoded", 0),
        ("attorney_validation_required", "all rows"),
        ("no_guarantee", "Controls and wording reduce operational risk; they cannot prevent or preclude litigation."),
    ]
    for row in method_rows:
        methodology.append(row)
    style_sheet(methodology)
    workbook.save(path)
    normalize_ooxml(path)


def draw_receipt_page(pdf: canvas.Canvas, scenario: dict[str, Any], receipt_no: int, mutation: str) -> None:
    width, height = letter
    pdf.setFillColor(colors.HexColor("#F4F7FA"))
    pdf.rect(0, 0, width, height, fill=1, stroke=0)
    left, top, receipt_width = 126, height - 55, 360
    pdf.setFillColor(colors.white)
    pdf.setStrokeColor(colors.HexColor("#B7C4D3"))
    pdf.roundRect(left, 52, receipt_width, height - 104, 8, fill=1, stroke=1)
    pdf.setFillColor(colors.HexColor("#17365D"))
    pdf.setFont("Helvetica-Bold", 15)
    pdf.drawCentredString(left + receipt_width / 2, top - 20, "RETAILGUARD MARKET")
    pdf.setFont("Helvetica", 8)
    pdf.drawCentredString(left + receipt_width / 2, top - 34, "SYNTHETIC TRAINING RECEIPT — NOT A REAL SALE")
    y = top - 58
    pdf.setFillColor(colors.black)
    pdf.setFont("Courier", 9)
    meta = [
        f"STORE {scenario['store_id']:<12} LANE {scenario['lane']}",
        f"TX {scenario['transaction_id']}-{receipt_no}",
        f"TIME {scenario['occurred_at']}",
        "-" * 47,
    ]
    for line in meta:
        pdf.drawString(left + 24, y, line)
        y -= 14
    lines = [
        ("Organic oranges bag", 1, scenario["duplicate_price"]),
        ("Organic oranges bag", 1, scenario["duplicate_price"] if mutation == "flagged" else 0),
        ("Sparkling water", 1, scenario["charged"] if mutation != "corrected" else scenario["displayed"]),
        (f"Produce wt {scenario['weight_charged']:.2f} lb", 1, round(2.99 * scenario["weight_charged"], 2)),
    ]
    total = 0.0
    for description, quantity, amount in lines:
        if amount <= 0:
            continue
        total += amount
        pdf.drawString(left + 24, y, description[:29])
        pdf.drawRightString(left + receipt_width - 24, y, f"{amount:>7.2f}")
        y -= 14
    tax = scenario["tax"]
    pdf.drawString(left + 24, y, "-" * 47)
    y -= 14
    for label, amount in (("SUBTOTAL", total), ("TAX", tax), ("TOTAL", total + tax)):
        pdf.setFont("Courier-Bold" if label == "TOTAL" else "Courier", 9)
        pdf.drawString(left + 24, y, label)
        pdf.drawRightString(left + receipt_width - 24, y, f"{amount:>7.2f}")
        y -= 16
    pdf.setFont("Courier", 8)
    y -= 5
    footer = [
        "Please review item, quantity, weight, price,",
        "discounts, and total. Ask about any difference.",
        "Keep this receipt for prompt price review.",
        "This notice does not limit statutory rights.",
    ]
    for line in footer:
        pdf.drawCentredString(left + receipt_width / 2, y, line)
        y -= 12
    pdf.setFillColor(colors.HexColor("#A61B1B") if mutation == "flagged" else colors.HexColor("#1F6B45"))
    pdf.setFont("Helvetica-Bold", 9)
    label = "EVIDENCE: FLAGGED TRANSACTION" if mutation == "flagged" else "CONTROL SAMPLE: CORRECTED TRANSACTION"
    pdf.drawCentredString(left + receipt_width / 2, 70, label)
    pdf.showPage()


def build_receipts(path: Path, scenario: dict[str, Any]) -> None:
    pdf = canvas.Canvas(str(path), pagesize=letter, pageCompression=1, invariant=1)
    pdf.setTitle(f"{scenario['scenario_id']} synthetic receipt evidence")
    pdf.setAuthor("RetailGuard benchmark generator")
    draw_receipt_page(pdf, scenario, 1, "flagged")
    draw_receipt_page(pdf, scenario, 2, "corrected")
    pdf.save()


def structure_signature(documents: Path) -> dict[str, Any]:
    signature: dict[str, Any] = {"files": sorted(path.name for path in documents.iterdir() if path.is_file())}
    signature["workbooks"] = {}
    for path in sorted(documents.glob("*.xlsx")):
        workbook = load_workbook(path, read_only=True, data_only=False)
        signature["workbooks"][path.name] = {
            "sheets": workbook.sheetnames,
            "columns": {sheet.title: sheet.max_column for sheet in workbook.worksheets},
        }
        workbook.close()
    signature["docx"] = {}
    for path in sorted(documents.glob("*.docx")):
        document = Document(path)
        signature["docx"][path.name] = [
            paragraph.text for paragraph in document.paragraphs
            if paragraph.style and paragraph.style.name.startswith("Heading")
        ]
    signature["pdf_pages"] = {path.name: len(PdfReader(path).pages) for path in documents.glob("*.pdf")}
    return signature


def seed_data() -> dict[str, list[dict[str, Any]]]:
    contract = json.loads(CONTRACT.read_text("utf-8"))
    rows_by_table = {table["name"]: [] for table in contract["tables"]}
    for index, scenario in enumerate(SCENARIOS, 1):
        rows_by_table["rc_incidents"].append({
            "id": index,
            "incident_number": scenario["incident_number"],
            "store_id": scenario["store_id"],
            "jurisdiction_code": scenario["jurisdiction_code"],
            "occurred_at": scenario["occurred_at"],
            "channel": "self_checkout",
            "allegation_type": "duplicate_scan_price_and_weight",
            "status": "open",
            "affected_transactions": 1,
            "estimated_exposure": round(
                scenario["duplicate_price"] + scenario["charged"] - scenario["displayed"]
                + (scenario["weight_charged"] - scenario["weight_label"]) * 2.99, 2
            ),
            "owner": "price-accuracy@retailguard.example",
            "root_cause": "pending validation",
            "remediation_status": "not_started",
        })
        rows_by_table["rc_transactions"].append({
            "id": index,
            "transaction_id": scenario["transaction_id"],
            "incident_id": index,
            "jurisdiction_code": scenario["jurisdiction_code"],
            "checkout_lane": scenario["lane"],
            "occurred_at": scenario["occurred_at"],
            "subtotal": round(2 * scenario["duplicate_price"] + scenario["charged"] + 2.99 * scenario["weight_charged"], 2),
            "tax": scenario["tax"],
            "total": round(2 * scenario["duplicate_price"] + scenario["charged"] + 2.99 * scenario["weight_charged"] + scenario["tax"], 2),
            "receipt_hash": f"synthetic-{scenario['scenario_id'].lower()}",
            "review_status": "flagged",
            "duplicate_scan_count": 1,
            "overcharge_total": rows_by_table["rc_incidents"][-1]["estimated_exposure"],
        })
        line_id = (index - 1) * 3
        rows_by_table["rc_receipt_lines"].extend([
            {"id": line_id + 1, "transaction_id": scenario["transaction_id"], "sku": "04120-88214", "description": "Organic oranges bag", "quantity": 2, "displayed_unit_price": scenario["duplicate_price"], "charged_unit_price": scenario["duplicate_price"], "label_weight": scenario["weight_label"], "charged_weight": scenario["weight_label"], "extended_price": scenario["duplicate_price"] * 2, "anomaly_type": "duplicate_scan"},
            {"id": line_id + 2, "transaction_id": scenario["transaction_id"], "sku": "00854-10219", "description": "Sparkling water", "quantity": 1, "displayed_unit_price": scenario["displayed"], "charged_unit_price": scenario["charged"], "label_weight": None, "charged_weight": None, "extended_price": scenario["charged"], "anomaly_type": "shelf_register_mismatch"},
            {"id": line_id + 3, "transaction_id": scenario["transaction_id"], "sku": "02941-55003", "description": "Random-weight produce", "quantity": 1, "displayed_unit_price": 2.99, "charged_unit_price": 2.99, "label_weight": scenario["weight_label"], "charged_weight": scenario["weight_charged"], "extended_price": round(2.99 * scenario["weight_charged"], 2), "anomaly_type": "weight_variance"},
        ])
    rows_by_table["rc_jurisdiction_rules"] = jurisdiction_rows()
    for table in contract["tables"]:
        if rows_by_table[table["name"]]:
            continue
        for index, row in enumerate((table.get("seed") or {}).get("rows") or [], 1):
            rows_by_table[table["name"]].append({"id": index, **row})
    # Three comparable failed audits give every mutation an executable record.
    rows_by_table["rc_price_audits"] = [
        {"id": index, "store_id": scenario["store_id"], "jurisdiction_code": scenario["jurisdiction_code"], "sample_size": 50, "accurate_items": 47, "accuracy_rate": 94.0, "standard": "NIST HB 130 EPPV research baseline; apply controlling local law", "passed": 0, "audit_date": "2026-08-19", "corrective_action": "open"}
        for index, scenario in enumerate(SCENARIOS, 1)
    ]
    return rows_by_table


def build(output_root: Path) -> dict[str, Any]:
    output_root.mkdir(parents=True, exist_ok=True)
    scenario_records = []
    signatures = []
    for scenario in SCENARIOS:
        scenario_root = output_root / scenario["slug"]
        documents = scenario_root / "documents"
        documents.mkdir(parents=True, exist_ok=True)
        build_incident_doc(documents / "incident-report.docx", scenario)
        build_event_log(documents / "checkout-event-log.xlsx", scenario)
        build_receipts(documents / "sample-receipts.pdf", scenario)
        build_policy_doc(documents / "customer-price-accuracy-policy.docx", scenario)
        build_jurisdiction_register(documents / "jurisdiction-source-register.xlsx", scenario)
        build_jurisdiction_register_v2(documents / "jurisdiction-authority-map-v2.xlsx", scenario)
        signature = structure_signature(documents)
        signatures.append(signature)
        files = [
            {"path": path.name, "bytes": path.stat().st_size, "sha256": sha256(path)}
            for path in sorted(documents.iterdir()) if path.is_file()
        ]
        record = {
            "schema_version": 2,
            "scenario_id": scenario["scenario_id"],
            "synthetic": True,
            "jurisdiction_code": scenario["jurisdiction_code"],
            "source_scenario": SCENARIOS[0]["scenario_id"] if scenario is not SCENARIOS[0] else None,
            "mutation_dimensions": [] if scenario is SCENARIOS[0] else ["jurisdiction", "facts", "amounts", "timestamps", "control_failure"],
            "structure_signature": signature,
            "files": files,
        }
        (scenario_root / "manifest.json").write_text(json.dumps(record, indent=2, sort_keys=True) + "\n", "utf-8")
        scenario_records.append(record)
    if any(signature != signatures[0] for signature in signatures[1:]):
        raise RuntimeError("mutated scenario structures do not match the base scenario")
    seeds = seed_data()
    (output_root / "seed-data.json").write_text(json.dumps({"schema_version": 1, "tables": seeds}, indent=2, sort_keys=True) + "\n", "utf-8")
    research_v2 = jurisdiction_research_v2_rows()
    report = {
        "schema_version": 2,
        "built_at": BUILD_DATE,
        "synthetic": True,
        "scenarios": len(SCENARIOS),
        "documents": sum(len(record["files"]) for record in scenario_records),
        "formats": {"docx": 6, "xlsx": 9, "pdf": 3},
        "jurisdictions": len(JURISDICTIONS),
        "primary_source_triaged_jurisdictions": sorted(PRIMARY_RULES),
        "pending_substantive_validation_jurisdictions": len(JURISDICTIONS) - len(PRIMARY_RULES),
        "authority_map_v2_jurisdictions": len(research_v2),
        "authority_map_v2_specific_authorities": sum(bool(row["citation"]) for row in research_v2),
        "authority_map_v2_substantive_legal_opinions": sum(bool(row["substantive_legal_opinion"]) for row in research_v2),
        "authority_map_v2_private_remedies_encoded": sum(bool(row["private_remedy_encoded"]) for row in research_v2),
        "authority_map_v2_all_attorney_validation_required": all(row["attorney_validation_required"] for row in research_v2),
        "authority_map_v2_source_kinds": {
            source_kind: sum(row["source_kind"] == source_kind for row in research_v2)
            for source_kind in sorted({row["source_kind"] for row in research_v2})
        },
        "all_attorney_validation_required": all(row["attorney_validation_required"] for row in seeds["rc_jurisdiction_rules"]),
        "shared_structure_signature": signatures[0],
        "scenario_manifests": [str((output_root / scenario["slug"] / "manifest.json").relative_to(ROOT)) for scenario in SCENARIOS],
        "seed_data": str((output_root / "seed-data.json").relative_to(ROOT)),
    }
    (output_root / "build-report.json").write_text(json.dumps(report, indent=2, sort_keys=True) + "\n", "utf-8")
    return report


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, default=OUTPUT_ROOT)
    parser.add_argument("--check", action="store_true", help="rebuild in isolation and require byte-identical artifacts")
    args = parser.parse_args()
    output = args.out.resolve()
    if args.check:
        (ROOT / "dist").mkdir(exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="retail-pack-check-", dir=ROOT / "dist") as temporary:
            rebuilt = Path(temporary) / "retail-price-accuracy"
            build(rebuilt)
            relative_paths = [
                path.relative_to(output)
                for path in output.glob("**/*")
                if path.is_file() and (
                    "documents" in path.parts
                    or path.name == "manifest.json" and "sources" not in path.parts
                    or path.name == "seed-data.json"
                )
            ]
            missing = [str(path) for path in relative_paths if not (rebuilt / path).is_file()]
            changed = [
                str(path) for path in relative_paths
                if (rebuilt / path).is_file()
                and sha256(output / path) != sha256(rebuilt / path)
            ]
            extras = [
                str(path.relative_to(rebuilt))
                for path in rebuilt.glob("**/*")
                if path.is_file()
                and (
                    "documents" in path.parts
                    or path.name == "manifest.json"
                    or path.name == "seed-data.json"
                )
                and path.relative_to(rebuilt) not in set(relative_paths)
            ]
            if missing or changed or extras:
                print(json.dumps({"missing": missing, "changed": changed, "extras": extras}, indent=2), file=sys.stderr)
                return 2
        print("retail price-accuracy pack is byte-reproducible")
        return 0
    print(json.dumps(build(output), sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
