from __future__ import annotations

import copy
import hashlib
import importlib.util
import json
import subprocess
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook


ROOT = Path(__file__).resolve().parents[2]


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


BROAD = load_module("lab_mutate_for_tests", ROOT / "research" / "lab_mutate.py")
STRICT = load_module("mutate_harvey_task_for_tests", ROOT / "tools" / "mutate_harvey_task.py")
AUDIT = load_module("audit_harvey_inputs_for_tests", ROOT / "tools" / "audit_harvey_inputs.py")
INGEST = load_module("lab_ingest_for_tests", ROOT / "world" / "ingest" / "lab_ingest.py")


def write_json(path: Path, value: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, indent=2) + "\n", encoding="utf-8")


def set_archive_comment(path: Path, comment: bytes) -> None:
    with zipfile.ZipFile(path, "a") as archive:
        archive.comment = comment


def tree_hashes(root: Path) -> dict[str, str]:
    return {
        path.relative_to(root).as_posix(): hashlib.sha256(path.read_bytes()).hexdigest()
        for path in root.rglob("*")
        if path.is_file()
    }


def rewrite_zip_copy(
    source: Path,
    destination: Path,
    member: str,
    old: bytes,
    new: bytes,
) -> None:
    replacements = 0
    with zipfile.ZipFile(source) as input_archive, zipfile.ZipFile(destination, "w") as output_archive:
        output_archive.comment = input_archive.comment
        for info in input_archive.infolist():
            payload = input_archive.read(info.filename)
            if info.filename == member:
                replacements = payload.count(old)
                payload = payload.replace(old, new)
            output_archive.writestr(info, payload)
    if replacements < 1:
        raise AssertionError(f"fixture token {old!r} missing from {member}")


class HarveyMutationFixture(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary = tempfile.TemporaryDirectory(prefix="harvey-mutation-tests-")
        self.addCleanup(self.temporary.cleanup)
        self.workspace = Path(self.temporary.name)
        self.source = self.workspace / "source"
        self.task_name = "banking-finance/source-task"
        self.task_dir = self.source / "tasks" / "banking-finance" / "source-task"
        documents = self.task_dir / "documents"
        documents.mkdir(parents=True)

        task = {
            "title": "Oregon Memo Review",
            "instructions": "Prepare a Memo about Oregon using the supplied evidence.",
            "deliverables": {"memo.docx": "memo.docx"},
            "criteria": [{
                "id": "C-001",
                "title": "Uses the Oregon record",
                "deliverables": ["memo.docx"],
                "match_criteria": "PASS if the Memo uses the Oregon record. FAIL if it does not.",
            }],
        }
        write_json(self.task_dir / "task.json", task)
        document_path = documents / "evidence.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Ore")
        paragraph.add_run("gon")
        document.save(document_path)
        set_archive_comment(document_path, b"fixture-comment")

        workbook_path = documents / "workbook.xlsx"
        workbook = Workbook()
        workbook.active.title = "Oregon"
        workbook.active["A1"] = "Oregon"
        workbook.save(workbook_path)
        set_archive_comment(workbook_path, b"fixture-comment")
        (documents / "Oregon-notes.txt").write_text(
            "Oregon evidence supports the Memo.\n", encoding="utf-8"
        )
        (self.source / "LICENSE").write_text("MIT fixture license\n", encoding="utf-8")
        self.entities = self.workspace / "entities.json"
        write_json(self.entities, [{"type": "place", "name": "Oregon"}])

        subprocess.run(["git", "init", "-q", str(self.source)], check=True)
        subprocess.run(["git", "-C", str(self.source), "config", "user.name", "Fixture"], check=True)
        subprocess.run(
            ["git", "-C", str(self.source), "config", "user.email", "fixture@example.test"],
            check=True,
        )
        subprocess.run(
            ["git", "-C", str(self.source), "remote", "add", "origin", "https://example.test/harvey-labs"],
            check=True,
        )
        subprocess.run(["git", "-C", str(self.source), "add", "."], check=True)
        subprocess.run(
            ["git", "-C", str(self.source), "commit", "-q", "-m", "fixture"], check=True
        )

        self.recipe = {
            "schema_version": 1,
            "variant_id": "memo-note-v1",
            "source_task": self.task_name,
            "output_task": "banking-finance/memo-note/v1",
            "layout_preserving": True,
            "mutation_axes": ["terminology"],
            "replacements": [{
                "old": "Memo",
                "new": "Note",
                "extensions": [".txt"],
                "min_document_occurrences": 1,
                "min_task_occurrences": 2,
            }],
        }

    def broad_generate(self, output_root: Path, seed: int = 3) -> Path:
        return BROAD.generate(
            self.task_name,
            self.entities,
            seed,
            self.source,
            output_root,
        )


class BroadMutationTests(HarveyMutationFixture):
    def test_seed_is_byte_deterministic_and_self_verifying(self) -> None:
        first_root = self.workspace / "first"
        second_root = self.workspace / "second"
        first = self.broad_generate(first_root)
        second = self.broad_generate(second_root)
        self.assertEqual(tree_hashes(first), tree_hashes(second))
        BROAD.check_generated(first, self.source, quiet=True)
        self.assertEqual(
            BROAD.package_topology(self.task_dir / "documents" / "workbook.xlsx"),
            BROAD.package_topology(first / "documents" / "workbook.xlsx"),
        )

    def test_split_run_mutation_preserves_run_structure(self) -> None:
        generated = self.broad_generate(self.workspace / "output")
        source_path = self.task_dir / "documents" / "evidence.docx"
        output_path = generated / "documents" / "evidence.docx"
        source_runs = Document(source_path).paragraphs[0].runs
        output_runs = Document(output_path).paragraphs[0].runs
        self.assertEqual(len(source_runs), len(output_runs))
        self.assertNotIn("Oregon", "".join(run.text for run in output_runs))
        self.assertEqual(
            BROAD.package_topology(source_path),
            BROAD.package_topology(output_path),
        )

    def test_seed_must_be_a_non_negative_integer(self) -> None:
        for seed in (-1, True, 1.5):
            with self.subTest(seed=seed), self.assertRaises(ValueError):
                BROAD.generate(
                    self.task_name,
                    self.entities,
                    seed,
                    self.source,
                    self.workspace / "output",
                )

    def test_existing_output_is_not_deleted(self) -> None:
        output_root = self.workspace / "output"
        generated = self.broad_generate(output_root)
        sentinel = generated / "keep-me"
        sentinel.write_text("preserve", encoding="utf-8")
        with self.assertRaises(FileExistsError):
            self.broad_generate(output_root)
        self.assertEqual(sentinel.read_text(encoding="utf-8"), "preserve")

    def test_task_path_traversal_is_rejected(self) -> None:
        with self.assertRaises(ValueError):
            BROAD.generate(
                "../outside",
                self.entities,
                1,
                self.source,
                self.workspace / "output",
            )

    def test_entities_symlink_is_rejected(self) -> None:
        link = self.workspace / "entities-link.json"
        try:
            link.symlink_to(self.entities)
        except OSError as error:
            self.skipTest(f"symlinks unavailable: {error}")
        with self.assertRaisesRegex(ValueError, "entities config is missing or unsafe"):
            BROAD.generate(
                self.task_name,
                link,
                1,
                self.source,
                self.workspace / "output",
            )

    def test_replacement_strings_are_not_regex_templates(self) -> None:
        self.assertEqual(BROAD.replace_text("Oregon", [("Oregon", r"A\1")]), r"A\1")

    def test_checker_detects_tampering(self) -> None:
        generated = self.broad_generate(self.workspace / "output")
        text_path = next((generated / "documents").glob("*-notes.txt"))
        text_path.write_text("tampered\n", encoding="utf-8")
        with self.assertRaises(ValueError):
            BROAD.check_generated(generated, self.source, quiet=True)

    def test_task_validation_rejects_unknown_deliverables(self) -> None:
        task = json.loads((self.task_dir / "task.json").read_text(encoding="utf-8"))
        task["criteria"][0]["deliverables"] = ["unknown.docx"]
        with self.assertRaisesRegex(ValueError, "unknown deliverable"):
            BROAD.validate_task(task)


class StrictMutationTests(HarveyMutationFixture):
    def recipe_path(self, recipe: dict | None = None) -> Path:
        path = self.workspace / "recipes" / "recipe.json"
        write_json(path, recipe or self.recipe)
        return path

    def test_layout_preservation_uses_utf8_and_xml_byte_lengths(self) -> None:
        bad_utf8 = copy.deepcopy(self.recipe)
        bad_utf8["replacements"][0].update({"old": "é", "new": "a"})
        with self.assertRaisesRegex(ValueError, "UTF-8 byte length"):
            STRICT.validate_recipe(bad_utf8)
        bad_xml = copy.deepcopy(self.recipe)
        bad_xml["replacements"][0].update({"old": "&", "new": "x"})
        with self.assertRaisesRegex(ValueError, "XML-escaped byte length"):
            STRICT.validate_recipe(bad_xml)

    def test_criterion_override_values_must_be_objects(self) -> None:
        invalid = copy.deepcopy(self.recipe)
        invalid["criterion_overrides"] = {"C-001": "not-an-object"}
        with self.assertRaisesRegex(ValueError, "values must be objects"):
            STRICT.validate_recipe(invalid)

    def test_recipe_symlink_is_rejected(self) -> None:
        recipe_path = self.recipe_path()
        link = self.workspace / "recipe-link.json"
        try:
            link.symlink_to(recipe_path)
        except OSError as error:
            self.skipTest(f"symlinks unavailable: {error}")
        with self.assertRaisesRegex(ValueError, "recipe is missing or unsafe"):
            STRICT.generate(link, self.source, self.workspace / "output")

    def test_generation_is_byte_deterministic_and_does_not_overwrite(self) -> None:
        recipe_path = self.recipe_path()
        first_root = self.workspace / "first-generated"
        second_root = self.workspace / "second-generated"
        first = STRICT.generate(recipe_path, self.source, first_root)
        second = STRICT.generate(recipe_path, self.source, second_root)
        self.assertEqual(tree_hashes(first), tree_hashes(second))
        sentinel = first / "keep-me"
        sentinel.write_text("preserve", encoding="utf-8")
        with self.assertRaises(FileExistsError):
            STRICT.generate(recipe_path, self.source, first_root)
        self.assertEqual(sentinel.read_text(encoding="utf-8"), "preserve")

    def test_ooxml_archive_comment_is_preserved(self) -> None:
        source = self.task_dir / "documents" / "workbook.xlsx"
        output = self.workspace / "mutated.xlsx"
        STRICT.mutate_ooxml(source, output, [{"old": "Oregon", "new": "Nevada"}])
        self.assertEqual(STRICT.package_topology(source), STRICT.package_topology(output))

    def test_split_run_residual_fails_closed(self) -> None:
        source = self.task_dir / "documents" / "evidence.docx"
        output = self.workspace / "mutated.docx"
        with self.assertRaisesRegex(ValueError, "source replacement text remains"):
            STRICT.copy_or_mutate(
                source,
                output,
                [{"old": "Oregon", "new": "Nevada", "extensions": [".docx"]}],
            )

    def test_invalid_generation_never_publishes(self) -> None:
        invalid = copy.deepcopy(self.recipe)
        invalid["replacements"][0]["min_document_occurrences"] = 99
        output_root = self.workspace / "generated"
        with self.assertRaisesRegex(ValueError, "expected at least 99"):
            STRICT.generate(self.recipe_path(invalid), self.source, output_root)
        self.assertFalse((output_root / "tasks" / "banking-finance" / "memo-note" / "v1").exists())

    def test_tree_checker_detects_orphans(self) -> None:
        recipes = self.workspace / "recipes"
        recipe_path = recipes / "recipe.json"
        write_json(recipe_path, self.recipe)
        output_root = self.workspace / "generated"
        STRICT.generate(recipe_path, self.source, output_root)
        STRICT.check_generated_root(output_root, recipes, self.source)
        orphan = output_root / "tasks" / "banking-finance" / "orphan"
        write_json(orphan / "task.json", {"orphan": True})
        with self.assertRaisesRegex(ValueError, "task set mismatch"):
            STRICT.check_generated_root(output_root, recipes, self.source)

    def test_checker_rederives_manifest_metadata(self) -> None:
        output_root = self.workspace / "generated"
        generated = STRICT.generate(self.recipe_path(), self.source, output_root)
        manifest_path = generated / "mutation-manifest.json"
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        manifest["files"][0]["source_bytes"] += 1
        write_json(manifest_path, manifest)
        with self.assertRaisesRegex(ValueError, "source byte count"):
            STRICT.check_generated(generated, self.source, quiet=True)


class AuditTests(HarveyMutationFixture):
    def test_small_corpus_audit_runs_in_process_pool(self) -> None:
        report_path = self.workspace / "audit-report.json"
        result = subprocess.run(
            [
                "python3",
                str(ROOT / "tools" / "audit_harvey_inputs.py"),
                "--source",
                str(self.source),
                "--report",
                str(report_path),
                "--workers",
                "2",
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, result.stderr)
        report = json.loads(report_path.read_text(encoding="utf-8"))
        self.assertEqual(report["schema_version"], 3)
        self.assertEqual(report["physical_inputs"], 3)
        self.assertEqual(report["format_validation"]["errors"], [])
        self.assertTrue(
            report["task_tree_coverage"]["all_tracked_task_files_classified"]
        )

    def test_ooxml_xml_parts_are_parsed(self) -> None:
        result = AUDIT.validate_input(self.task_dir / "documents" / "evidence.docx")
        self.assertGreaterEqual(result["xml_parts"], 3)
        self.assertEqual(result["xml_parts_failed"], 0)
        self.assertEqual(result["defects"], [])

    def test_unescaped_ampersand_is_classified_as_a_source_defect(self) -> None:
        path = self.workspace / "malformed.xlsx"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr(
                "[Content_Types].xml",
                '<?xml version="1.0"?><Types xmlns="urn:test"/>',
            )
            archive.writestr(
                "_rels/.rels",
                '<?xml version="1.0"?><Relationships xmlns="urn:test"/>',
            )
            archive.writestr(
                "xl/workbook.xml",
                '<?xml version="1.0"?><workbook xmlns="urn:test"/>',
            )
            archive.writestr(
                "xl/worksheets/sheet1.xml",
                '<?xml version="1.0"?><worksheet><t>A & B</t></worksheet>',
            )
        result = AUDIT.validate_input(path)
        self.assertEqual(result["xml_parts_failed"], 1)
        self.assertEqual(result["defects"][0]["kind"], "unescaped_ampersand")
        self.assertEqual(result["defects"][0]["occurrences"], 1)

    def test_symlink_input_is_rejected(self) -> None:
        link = self.workspace / "linked.docx"
        try:
            link.symlink_to(self.task_dir / "documents" / "evidence.docx")
        except OSError as error:
            self.skipTest(f"symlinks unavailable: {error}")
        with self.assertRaisesRegex(ValueError, "symlink"):
            AUDIT.validate_input(link)

    def test_nested_symlink_directory_is_reported(self) -> None:
        documents = self.task_dir / "documents"
        target = self.workspace / "linked-directory-target"
        target.mkdir()
        (target / "hidden.txt").write_text("hidden\n", encoding="utf-8")
        link = documents / "linked-directory"
        try:
            link.symlink_to(target, target_is_directory=True)
        except OSError as error:
            self.skipTest(f"symlinks unavailable: {error}")
        files, errors = AUDIT.inventory_document_set(documents, self.source)
        self.assertTrue(any("symlink" in error for error in errors))
        self.assertNotIn(link / "hidden.txt", files)

    def test_known_defect_allowlist_symlink_is_rejected(self) -> None:
        allowlist = self.workspace / "known-defects.json"
        write_json(allowlist, {
            "schema_version": 1,
            "source_repo": "harveyai/harvey-labs",
            "source_commit": "0000000",
            "defects": [],
        })
        link = self.workspace / "known-defects-link.json"
        try:
            link.symlink_to(allowlist)
        except OSError as error:
            self.skipTest(f"symlinks unavailable: {error}")
        with self.assertRaisesRegex(ValueError, "missing or unsafe"):
            AUDIT.load_expected_defects(link, "0000000")


class IngestRecoveryTests(HarveyMutationFixture):
    def test_docx_raw_ampersand_is_recovered_without_source_rewrite(self) -> None:
        source = self.task_dir / "documents" / "evidence.docx"
        malformed = self.workspace / "malformed.docx"
        rewrite_zip_copy(
            source,
            malformed,
            "word/document.xml",
            b">Ore<",
            b">A &<",
        )
        before = malformed.read_bytes()
        text, recovery = INGEST.parse_document_with_recovery(malformed, ".docx")
        self.assertIn("A &gon", text)
        self.assertEqual(recovery["occurrences"], 1)
        self.assertEqual(malformed.read_bytes(), before)

    def test_xlsx_recovery_metadata_survives_cached_parse(self) -> None:
        source = self.task_dir / "documents" / "workbook.xlsx"
        malformed = self.workspace / "malformed.xlsx"
        rewrite_zip_copy(
            source,
            malformed,
            "xl/worksheets/sheet1.xml",
            b">Oregon<",
            b">A & B<",
        )
        destination = self.workspace / "store"
        first = INGEST.parse_one(
            (str(self.workspace), str(destination), malformed.name, ".xlsx")
        )
        second = INGEST.parse_one(
            (str(self.workspace), str(destination), malformed.name, ".xlsx")
        )
        self.assertEqual(first["status"], "parsed")
        self.assertEqual(first["recovery"], second["recovery"])
        self.assertEqual(first["recovery"]["occurrences"], 1)


if __name__ == "__main__":
    unittest.main()
